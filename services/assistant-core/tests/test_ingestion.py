from __future__ import annotations

import io

import pytest

from violet_assistant.ingestion.extractors import (
    ExtractionError,
    extract_text,
    is_image,
)
from violet_assistant.ingestion.ocr import mime_for


def test_extract_csv_sniffs_and_summarizes() -> None:
    data = b"name,score\nAda,90\nBeb,75\n"
    result = extract_text("scores.csv", data)
    assert result["kind"] == "csv"
    assert "name | score" in result["text"]
    assert "Ada | 90" in result["text"]
    assert "2 cols" in result["text"]


def test_extract_plain_text_and_json() -> None:
    assert extract_text("note.md", b"# Title\nbody")["kind"] == "text"
    j = extract_text("d.json", b'{"b":2,"a":1}')
    assert j["kind"] == "json"
    assert '"b": 2' in j["text"]


def test_extract_docx() -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Meeting recap")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Task"
    table.rows[0].cells[1].text = "Owner"
    buf = io.BytesIO()
    doc.save(buf)

    result = extract_text("m.docx", buf.getvalue())
    assert result["kind"] == "docx"
    assert "Meeting recap" in result["text"]
    assert "Task | Owner" in result["text"]


def test_extract_xlsx() -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["a", "b"])
    ws.append([1, 2])
    buf = io.BytesIO()
    wb.save(buf)

    result = extract_text("s.xlsx", buf.getvalue())
    assert result["kind"] == "xlsx"
    assert "# Sheet: Data" in result["text"]
    assert "1 | 2" in result["text"]


def test_image_and_unsupported_raise() -> None:
    assert is_image("scan.png")
    with pytest.raises(ExtractionError):
        extract_text("scan.png", b"\x89PNG")
    with pytest.raises(ExtractionError):
        extract_text("thing.xyz", b"data")


def test_truncation_flag() -> None:
    big = ("word " * 6000).encode("utf-8")  # ~30k chars > 20k cap
    result = extract_text("big.txt", big)
    assert result["truncated"] is True
    assert "[truncated]" in result["text"]


def test_mime_for() -> None:
    assert mime_for("a.png") == "image/png"
    assert mime_for("a.JPG") == "image/jpeg"
    assert mime_for("a.webp") == "image/webp"
