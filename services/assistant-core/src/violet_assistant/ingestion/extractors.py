from __future__ import annotations

import csv
import io
import json
from pathlib import PurePosixPath


MAX_TEXT_CHARS = 20_000
TEXT_EXTS = {".txt", ".md", ".markdown", ".log", ".rst"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}


class ExtractionError(Exception):
    pass


def extension(filename: str) -> str:
    return PurePosixPath(filename).suffix.lower()


def is_image(filename: str) -> bool:
    return extension(filename) in IMAGE_EXTS


def _clip(text: str) -> tuple[str, bool]:
    text = text.strip()
    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS].rstrip() + "\n… [truncated]", True
    return text, False


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_csv(data: bytes, delimiter: str | None = None) -> str:
    text = _decode(data)
    sample = text[:4096]
    if delimiter is None:
        try:
            delimiter = csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
        except csv.Error:
            delimiter = ","
    rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    lines = [" | ".join(cell.strip() for cell in row) for row in rows if any(row)]
    header = f"[CSV: {len(rows)} rows x {len(rows[0]) if rows else 0} cols]\n" if rows else ""
    return header + "\n".join(lines)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
    workbook.close()
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"--- page {index} ---\n{text}")
    if not pages:
        raise ExtractionError(
            "No extractable text (the PDF may be scanned). Upload the page as an image for OCR."
        )
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(filename: str, data: bytes) -> dict:
    """Extract text from a document. Returns {kind, text, chars, truncated}.

    Raises ExtractionError for images (route to OCR) and unsupported/empty files.
    """
    ext = extension(filename)
    if ext in IMAGE_EXTS:
        raise ExtractionError("Image file — use OCR, not text extraction.")

    if ext == ".csv":
        kind, raw = "csv", _extract_csv(data)
    elif ext == ".tsv":
        kind, raw = "tsv", _extract_csv(data, delimiter="\t")
    elif ext in {".xlsx", ".xlsm"}:
        kind, raw = "xlsx", _extract_xlsx(data)
    elif ext == ".pdf":
        kind, raw = "pdf", _extract_pdf(data)
    elif ext == ".docx":
        kind, raw = "docx", _extract_docx(data)
    elif ext == ".json":
        try:
            kind, raw = "json", json.dumps(json.loads(_decode(data)), indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            kind, raw = "text", _decode(data)
    elif ext in TEXT_EXTS or ext == "":
        kind, raw = "text", _decode(data)
    else:
        raise ExtractionError(f"Unsupported file type: {ext or '(none)'}")

    text, truncated = _clip(raw)
    if not text:
        raise ExtractionError("File contained no readable text.")
    return {"kind": kind, "text": text, "chars": len(text), "truncated": truncated}
